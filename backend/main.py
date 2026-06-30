import os
import shutil
import html
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, status, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv

# 1. Google GenAI SDK
from google import genai
from google.genai import types

# ReportLab (PDF)
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# python-docx (Word)
from docx import Document

# Custom local imports
try:
    from resume_parser import extract_text_from_pdf
    from skills import extract_skills
    from ats_score import calculate_ats_score
except ImportError as e:
    print(f"Warning: Local modules not found. Essential features might fail. Error: {e}")

# CONFIGURATION & INITIALIZATION
load_dotenv()

app = FastAPI(title="AI Resume Analyzer API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Gemini Client Initialization
client = genai.Client()

UPLOAD_DIR = "uploads"
TEMP_DIR = "temp_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)


# Pydantic Models
class ChatRequest(BaseModel):
    message: str

class ResumeData(BaseModel):
    fullName: str
    email: str
    phone: str
    summary: str
    experience: str
    education: str
    skills: str
    certificates: str
    achievements: str
    extracurricular: str
    hobbies: str
    languages: str

class SRSData(BaseModel):
    project_title: str
    introduction: str
    objectives: list[str]
    scope: str
    functional_requirements: list[str]
    non_functional_requirements: list[str]
    user_roles: list[str]
    modules: list[str]
    database_tables: list[str]
    tech_stack: dict
    timeline: list


# Helper function to remove temp files after response is sent
def remove_file(path: str):
    if os.path.exists(path):
        try:
            os.remove(path)
        except Exception as e:
            print(f"Error removing temp file {path}: {e}")


@app.get("/")
def home():
    return {"message": "AI Resume Analyzer API Running"}


# GEMINI FEEDBACK FUNCTION
def get_ai_feedback(resume_text: str, job_description: str) -> str:
    prompt = f"""
    Analyze this resume against the provided job description.

    Resume:
    {resume_text}

    Job Description:
    {job_description}

    Provide the following structured details:
    1. Resume strengths
    2. Missing skills
    3. Improvement suggestions
    4. ATS improvement tips
    5. Career advice
    """

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.5
        )
    )
    return response.text


# ANALYZE ROUTE
@app.post("/analyze")
def analyze_resume(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    
    # Save Uploaded File
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    finally:
        file.file.close() # ஃபைல் மெமரியை சரியாக மூடுதல்

    try:
        # Extract Text & Metrics
        resume_text = extract_text_from_pdf(file_path)
        skills = extract_skills(resume_text)
        ats_score = calculate_ats_score(resume_text, job_description)
        ai_feedback = get_ai_feedback(resume_text, job_description)

        # Dynamic Missing Skills Check
        missing_skills = []
        target_skills = {"python", "react", "sql", "docker", "aws", "tensorflow", "fastapi"}
        
        cleaned_jd_text = job_description.lower().translate(str.maketrans("", "", ",.?!;:()"))
        jd_words = set(cleaned_jd_text.split())

        for word in target_skills:
            if word in jd_words and word not in [s.lower() for s in skills]:
                missing_skills.append(word)

        # Rule-Based Suggestions
        suggestions = []
        resume_lower = resume_text.lower()

        if ats_score < 50:
            suggestions.append("Improve ATS score by adding more keywords tailored to the job profile.")
        if len(skills) < 5:
            suggestions.append("Expand your technical skills section to showcase your full toolkit.")
        if "project" not in resume_lower:
            suggestions.append("Add a projects section featuring your best work and GitHub links.")
        if "experience" not in resume_lower:
            suggestions.append("Add professional experience, internship, or open-source contribution details.")
        if "docker" not in resume_lower and "docker" in cleaned_jd_text:
            suggestions.append("Consider learning and adding Docker for modern backend/containerized roles.")
        if "aws" not in resume_lower and "aws" in cleaned_jd_text:
            suggestions.append("Add cloud computing infrastructure knowledge such as AWS.")

        return {
            "filename": file.filename,
            "skills_found": skills,
            "ats_score": ats_score,
            "missing_skills": missing_skills,
            "suggestions": suggestions,
            "ai_feedback": ai_feedback,
            "resume_preview": resume_text[:1000]
        }
        
    except Exception as e:
        print(f"Error during analysis: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An error occurred while processing the resume: {str(e)}"
        )
        
    finally:
        # ஃபைலை நீக்குதல்
        if os.path.exists(file_path):
            os.remove(file_path)


# CHATBOT ASSISTANT ROUTE
@app.post("/Chatbot")
def chatbot_assistant(request: ChatRequest):
    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=request.message,
            config=types.GenerateContentConfig(
                system_instruction=(
                    "You are an expert AI Career and Resume Assistant. Help the user with resume tips, "
                    "interview preparation, skill development, and career growth. Keep answers concise and helpful."
                ),
                temperature=0.7,
            )
        )
        return {"reply": response.text}
    except Exception as e:
        print(f"Chatbot Error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Sorry, I am facing some connection issues. Please try again later."
        )


# PDF GENERATION ENDPOINT
@app.post("/api/resume/download-pdf")
def generate_pdf(data: ResumeData, background_tasks: BackgroundTasks):
    # திருத்தம் 1: பெயர் காலியாக இருந்தால் 'Resume' என மாற்றுதல்
    clean_name = "".join(c for c in data.fullName if c.isalnum() or c in (' ', '_', '-')).rstrip()
    safe_filename = clean_name if clean_name else "Resume"
    pdf_path = os.path.join(TEMP_DIR, f"{safe_filename}.pdf")
    
    try:
        doc = SimpleDocTemplate(pdf_path, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
        story = []
        styles = getSampleStyleSheet()
        
        name_style = ParagraphStyle(
            'NameStyle', parent=styles['Heading1'], fontSize=24, leading=28,
            textColor=colors.HexColor('#00E5FF'), spaceAfter=4
        )
        contact_style = ParagraphStyle(
            'ContactStyle', parent=styles['Normal'], fontSize=10, leading=14,
            textColor=colors.HexColor('#52647c'), spaceAfter=15
        )
        h2_style = ParagraphStyle(
            'H2Style', parent=styles['Heading2'], fontSize=14, leading=18,
            textColor=colors.HexColor('#FF4FD8'), spaceBefore=12, spaceAfter=6
        )
        body_style = ParagraphStyle(
            'BodyStyle', parent=styles['BodyText'], fontSize=11, leading=16,
            textColor=colors.HexColor('#111111'), alignment=4
        )

        # திருத்தம் 2: XML எர்ரர்களைத் தவிர்க்க html.escape() செய்தல்
        def format_text(text: str) -> str:
            if not text:
                return "Not Provided"
            escaped = html.escape(text)
            return escaped.replace('\n', '<br/>')

        story.append(Paragraph(format_text(data.fullName).upper(), name_style))
        story.append(Paragraph(f"Email: {html.escape(data.email or 'N/A')}  |  Phone: {html.escape(data.phone or 'N/A')}", contact_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#dddddd'), spaceAfter=15))

        sections = [
            ("PROFESSIONAL SUMMARY", data.summary),
            ("WORK EXPERIENCE", data.experience)
        ]
        for title, content in sections:
            story.append(Paragraph(title, h2_style))
            story.append(Paragraph(format_text(content), body_style))
            story.append(Spacer(1, 10))

        h2_cyan_style = ParagraphStyle('H2Cyan', parent=h2_style, textColor=colors.HexColor('#00E5FF'))
        
        cyan_sections = [
            ("EDUCATION", data.education),
            ("SKILLS", data.skills)
        ]
        for title, content in cyan_sections:
            story.append(Paragraph(title, h2_cyan_style))
            story.append(Paragraph(format_text(content), body_style))
            story.append(Spacer(1, 10))

        second_cyan_sections = [
            ("CERTIFICATES", data.certificates),
            ("ACHIEVEMENTS", data.achievements),
            ("EXTRACURRICULAR ACTIVITIES", data.extracurricular),
            ("HOBBIES", data.hobbies),
            ("LANGUAGES", data.languages)
        ]
        for title, content in second_cyan_sections:
            story.append(Paragraph(title, h2_cyan_style))
            story.append(Paragraph(format_text(content), body_style))
            story.append(Spacer(1, 10))

        doc.build(story)
        
        background_tasks.add_task(remove_file, pdf_path)
        return FileResponse(pdf_path, media_type="application/pdf", filename=f"{safe_filename}.pdf")
    
    except Exception as e:
        # ஃபைல் கிராஷ் ஆனால் அதையும் நீக்கிவிட வேண்டும்
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        raise HTTPException(status_code=500, detail=f"PDF Generation Error: {str(e)}")


# WORD (DOCX) GENERATION ENDPOINT
@app.post("/api/resume/download-word")
def generate_word(data: ResumeData, background_tasks: BackgroundTasks):
    clean_name = "".join(c for c in data.fullName if c.isalnum() or c in (' ', '_', '-')).rstrip()
    safe_filename = clean_name if clean_name else "Resume"
    word_path = os.path.join(TEMP_DIR, f"{safe_filename}.docx")
    
    try:
        doc = Document()
        heading = doc.add_heading(level=1)
        run = heading.add_run((data.fullName or "YOUR NAME").upper())
        run.font.size = 24
        
        contact = doc.add_paragraph()
        contact.add_run(f"Email: {data.email or 'N/A'} | Phone: {data.phone or 'N/A'}")
        doc.add_paragraph("__________________________________________________________________")

        def add_resume_section(title, content):
            doc.add_heading(title, level=2)
            doc.add_paragraph(content or "Not Provided")
            doc.add_paragraph()

        add_resume_section("PROFESSIONAL SUMMARY", data.summary)
        add_resume_section("WORK EXPERIENCE", data.experience)
        add_resume_section("EDUCATION", data.education)
        add_resume_section("SKILLS", data.skills)
        add_resume_section("CERTIFICATES", data.certificates)
        add_resume_section("ACHIEVEMENTS", data.achievements)
        add_resume_section("EXTRACURRICULAR ACTIVITIES", data.extracurricular)
        add_resume_section("HOBBIES", data.hobbies)
        add_resume_section("LANGUAGES", data.languages)

        doc.save(word_path)
        
        background_tasks.add_task(remove_file, word_path)
        return FileResponse(
            word_path, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            filename=f"{safe_filename}.docx"
        )
        
    except Exception as e:
        if os.path.exists(word_path):
            os.remove(word_path)
        raise HTTPException(status_code=500, detail=f"Word Generation Error: {str(e)}")
class SRSRequest(BaseModel):
    project_title: str
    project_type: str


@app.post("/generate-srs")
async def generate_srs(data: SRSRequest):

    prompt = f"""
You are an expert Software Architect.

Generate a complete Software Requirements Specification (SRS).

Project Title:
{data.project_title}

Project Type:
{data.project_type}

Return ONLY valid JSON.

Structure:

{{
"project_title":"",
"introduction":"",
"objectives":["","",""],
"scope":"",
"functional_requirements":["","",""],
"non_functional_requirements":["","",""],
"user_roles":["","",""],
"modules":["","",""],
"database_tables":["","",""],
"tech_stack":{{
"frontend":"",
"backend":"",
"database":"",
"authentication":"",
"deployment":""
}},
"timeline":[
{{"week":"Week 1","task":""}},
{{"week":"Week 2","task":""}},
{{"week":"Week 3","task":""}},
{{"week":"Week 4","task":""}},
{{"week":"Week 5","task":""}},
{{"week":"Week 6","task":""}}
]
}}

Return JSON only.
"""

    try:

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.5
            )
        )

        text = response.text.strip()

        if text.startswith("```json"):
            text = text.replace("```json", "").replace("```", "").strip()
        elif text.startswith("```"):
            text = text.replace("```", "").strip()

        import json

        return json.loads(text)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/download-srs-docx")
def download_srs_docx(data:SRSData,background_tasks:BackgroundTasks):

    filename=f"{data.project_title.replace(' ','_')}.docx"

    docx_path=os.path.join(TEMP_DIR,filename)

    doc=Document()

    doc.add_heading(data.project_title,1)

    def section(title,content):

        doc.add_heading(title,2)

        if isinstance(content,list):

            for item in content:
                doc.add_paragraph(item,style="List Bullet")

        elif isinstance(content,dict):

            for k,v in content.items():
                doc.add_paragraph(f"{k.title()} : {v}")

        else:

            doc.add_paragraph(str(content))

    section("Introduction",data.introduction)
    section("Objectives",data.objectives)
    section("Scope",data.scope)
    section("Functional Requirements",data.functional_requirements)
    section("Non Functional Requirements",data.non_functional_requirements)
    section("User Roles",data.user_roles)
    section("Modules",data.modules)
    section("Database Tables",data.database_tables)
    section("Tech Stack",data.tech_stack)

    doc.add_heading("Development Timeline",2)

    for item in data.timeline:
        doc.add_paragraph(
            f"{item['week']} : {item['task']}"
        )

    doc.save(docx_path)

    background_tasks.add_task(remove_file,docx_path)

    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

@app.post("/download-srs-pdf")
def download_srs_pdf(data: SRSData, background_tasks: BackgroundTasks):

    filename = f"{data.project_title.replace(' ','_')}.pdf"

    pdf_path = os.path.join(TEMP_DIR, filename)

    doc = SimpleDocTemplate(pdf_path, pagesize=letter)

    styles = getSampleStyleSheet()

    story = []

    story.append(Paragraph(f"<b>{data.project_title}</b>", styles["Title"]))
    story.append(Spacer(1,12))

    def section(title, content):

        story.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        story.append(Spacer(1,5))

        if isinstance(content,list):

            for item in content:
                story.append(Paragraph(f"• {item}", styles["BodyText"]))

        elif isinstance(content,dict):

            for k,v in content.items():
                story.append(Paragraph(f"<b>{k.title()}</b>: {v}", styles["BodyText"]))

        else:

            story.append(Paragraph(str(content), styles["BodyText"]))

        story.append(Spacer(1,10))

    section("Introduction",data.introduction)
    section("Objectives",data.objectives)
    section("Scope",data.scope)
    section("Functional Requirements",data.functional_requirements)
    section("Non Functional Requirements",data.non_functional_requirements)
    section("User Roles",data.user_roles)
    section("Modules",data.modules)
    section("Database Tables",data.database_tables)
    section("Tech Stack",data.tech_stack)

    story.append(Paragraph("<b>Development Timeline</b>",styles["Heading2"]))

    for item in data.timeline:
        story.append(
            Paragraph(
                f"{item['week']} : {item['task']}",
                styles["BodyText"]
            )
        )

    doc.build(story)

    background_tasks.add_task(remove_file,pdf_path)

    return FileResponse(
        pdf_path,
        media_type="application/pdf",
        filename=filename
    )

@app.get("/api/roles")
def get_roles():
    # React-ன் roles.map() வேலை செய்ய இது ஒரு Array-ஆக (List) இருக்க வேண்டும்
    return ["Full Stack Developer", "Backend Engineer", "Frontend Developer", "Data Scientist"]

@app.get("/api/questions")
def get_questions(role: str = None):
    # React கேட்கும் மற்றொரு எண்ட்பாயிண்ட்
    return {
        "role": role,
        "questions": [
            "What is FastAPI?",
            "Explain React state management.",
            "How do you handle CORS in backend?"
        ]
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True)